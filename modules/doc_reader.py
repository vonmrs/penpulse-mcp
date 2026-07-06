"""
模块三：Word 文档读取 (doc_reader)
支持 .docx 文件读取，转换为 Markdown
"""

import os
import re
from typing import Dict


def read_docx(file_path: str) -> dict:
    """
    读取 .docx 文件，转换为 Markdown

    Args:
        file_path: .docx 文件路径

    Returns:
        dict with status, markdown content, title
    """
    if not os.path.exists(file_path):
        return {"status": "error", "message": f"文件不存在: {file_path}", "markdown": "", "title": ""}

    if not file_path.lower().endswith(".docx"):
        return {"status": "error", "message": "仅支持 .docx 格式文件", "markdown": "", "title": ""}

    try:
        from docx import Document
    except ImportError:
        return {
            "status": "error",
            "message": "缺少 python-docx 库，请运行: pip install python-docx",
            "markdown": "",
            "title": ""
        }

    try:
        doc = Document(file_path)
        markdown_parts = []
        title = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                markdown_parts.append("")
                continue

            style_name = para.style.name.lower()

            # 标题层级判断
            if "heading 1" in style_name or style_name == "title":
                text = f"# {text}"
                if not title:
                    title = text.replace("# ", "")
            elif "heading 2" in style_name:
                text = f"## {text}"
            elif "heading 3" in style_name:
                text = f"### {text}"
            elif "bold" in style_name or para.runs and para.runs[0].bold:
                # 粗体段落作为小标题
                text = f"**{text}**"
            else:
                # 清理格式残留
                text = _clean_text(text)

            markdown_parts.append(text)

        # 处理表格
        for table in doc.tables:
            md_table = _table_to_markdown(table)
            markdown_parts.append(md_table)

        markdown = "\n\n".join(markdown_parts)
        # 清理多余空行
        markdown = re.sub(r"\n{3,}", "\n\n", markdown)

        return {
            "status": "ok",
            "message": f"文档读取成功，共 {len(doc.paragraphs)} 段落",
            "markdown": markdown.strip(),
            "title": title or os.path.splitext(os.path.basename(file_path))[0]
        }

    except Exception as e:
        return {"status": "error", "message": f"读取失败: {str(e)}", "markdown": "", "title": ""}


def _clean_text(text: str) -> str:
    """清理 Word 文本中的格式残留"""
    # 去除多余空格
    text = re.sub(r"[ \t]+", " ", text)
    # 去除特殊空白字符
    text = re.sub(r"[\u200b-\u200f\ufeff]", "", text)
    # 保留基本标点，其他不变
    return text


def _table_to_markdown(table) -> str:
    """将 python-docx 表格转为 Markdown 表格"""
    rows = table.rows
    if not rows:
        return ""

    cols = len(rows[0].cells)
    lines = []

    # 表头
    header = [cell.text.strip() for cell in rows[0].cells]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * cols) + " |")

    # 数据行
    for row in rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)
